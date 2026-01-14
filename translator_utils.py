import pdfplumber
from deep_translator import GoogleTranslator
from docx import Document
import io

def extract_text_from_pdf(file):
    """
    Extracts text from a PDF file using pdfplumber, preserving basic structure 
    based on font size (e.g. headings vs body).
    Returns a list of dictionaries: {'text': str, 'size': float}
    """
    structured_content = []
    
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            # pdfplumber extract_words gives us x0, top, x1, bottom, text, size, etc.
            words = page.extract_words(keep_blank_chars=True, use_text_flow=True)
            
            # Simple heuristic: Group words into lines/blocks based on "top" position
            # This is a basic implementation. Ideally we'd cluster by Y position.
            if not words:
                continue

            current_line_y = words[0]['top']
            current_line_text = []
            current_font_size = words[0].get('size', 12) # Default if missing
            max_size_in_line = 0

            for word in words:
                # If word is on a new line (significant Y difference)
                if abs(word['top'] - current_line_y) > 5:  # Tolerance of 5px
                    # Commit previous line
                    full_line_text = " ".join(current_line_text).strip()
                    if full_line_text:
                        structured_content.append({
                            'text': full_line_text,
                            'size': max_size_in_line
                        })
                    
                    # Reset for new line
                    current_line_y = word['top']
                    current_line_text = [word['text']]
                    max_size_in_line = word.get('size', 0)
                else:
                    current_line_text.append(word['text'])
                    max_size_in_line = max(max_size_in_line, word.get('size', 0))
            
            # Commit last line
            full_line_text = " ".join(current_line_text).strip()
            if full_line_text:
                structured_content.append({
                    'text': full_line_text,
                    'size': max_size_in_line
                })
                
    return structured_content

def translate_content(content_blocks, source_lang='auto', target_lang='id'):
    """
    Translates a list of content blocks using batching to improve speed.
    """
    translator = GoogleTranslator(source=source_lang, target=target_lang)
    translated_blocks = []
    
    # Prepare batches
    batches = []
    current_batch_text = []
    current_batch_indices = []
    current_char_count = 0
    
    # Separator unlikely to be in text
    SEPARATOR = "\n|||\n" 
    
    for i, block in enumerate(content_blocks):
        text = block['text'].strip()
        if not text:
            # Keep empty blocks to preserve spacing/structure logic if needed, 
            # or just ignore. Let's append as is.
            continue
            
        # If adding this text exceeds limit (e.g. 4500 chars safe limit for Google free)
        # or just reasonable chunk
        if current_char_count + len(text) > 4000:
            # Commit current batch
            batches.append({
                'text': SEPARATOR.join(current_batch_text),
                'indices': current_batch_indices
            })
            # Reset
            current_batch_text = [text]
            current_batch_indices = [i]
            current_char_count = len(text)
        else:
            current_batch_text.append(text)
            current_batch_indices.append(i)
            current_char_count += len(text) + len(SEPARATOR)
            
    # Commit last batch
    if current_batch_text:
        batches.append({
            'text': SEPARATOR.join(current_batch_text),
            'indices': current_batch_indices
        })
    
    # Initialize result list with original content (to fill in blanks if any)
    # We will overwrite the 'text' field for successful translations
    results = [b.copy() for b in content_blocks]
    
    # Process batches
    for batch in batches:
        try:
            full_text = batch['text']
            indices = batch['indices']
            
            # Translate big chunk
            translated_chunk = translator.translate(full_text)
            
            # Split back
            # We use strict matching on separator. 
            # Sometimes ML translaters eat the separator or change it (e.g. "|||" -> "|||")
            # We need to be robust. 
            
            # Try splitting by the separator (stripped or not)
            parts = translated_chunk.split("|||")
            
            # Cleanup parts
            parts = [p.strip() for p in parts]
            
            # Safety check: if counts don't match, we might have a drift.
            # If so, we fallback to one-by-one for this specific batch or just map best effort.
            if len(parts) == len(indices):
                for idx, part in zip(indices, parts):
                    results[idx]['text'] = part
            else:
                # Fallback: Translate individually for this failed batch
                # preventing the whole doc from being messed up
                for idx in indices:
                    try:
                        results[idx]['text'] = translator.translate(content_blocks[idx]['text'])
                    except:
                        pass # keep original
        except Exception as e:
            # On fatal error for batch, try individual fallback
            indices = batch['indices']
            for idx in indices:
                try:
                    results[idx]['text'] = translator.translate(content_blocks[idx]['text'])
                except:
                    pass

    return results

def create_docx(translated_blocks):
    """
    Creates a Word document from the translated structured content.
    Uses font size to determine Heading levels.
    """
    doc = Document()
    
    # Calculate mode font size to determine "Body" text size
    sizes = [b['size'] for b in translated_blocks]
    if sizes:
        mode_size = max(set(sizes), key=sizes.count)
    else:
        mode_size = 12

    for block in translated_blocks:
        text = block['text']
        size = block['size']
        
        # Heuristic for Headings
        # If size is significantly larger than body text (e.g. +2pts)
        if size >= mode_size + 2:
            try:
                # Map size to heading level roughly
                # Bigger size -> Smaller Heading Level (H1 is biggest)
                diff = size - mode_size
                if diff > 10:
                    doc.add_heading(text, level=1)
                elif diff > 5:
                    doc.add_heading(text, level=2)
                else:
                    doc.add_heading(text, level=3)
            except:
                doc.add_paragraph(text)
        else:
            doc.add_paragraph(text)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
